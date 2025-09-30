import os
import tiktoken
from pathlib import Path
import random
from docx import Document
import re
from typing import List, Dict, Tuple
import pymupdf4llm

class RecursiveChunker:
    
    def __init__(self, 
                 max_tokens: int = 1000,
                 overlap_tokens: int = 100,
                 min_chunk_tokens: int = 200,
                 max_table_tokens: int = 600,  # New: Maximum size for preserved tables
                 target_utilization: float = 0.85,
                 encoding: str = 'cl100k_base'):
        self.max_tokens = max_tokens
        self.overlap_tokens = overlap_tokens
        self.min_chunk_tokens = min_chunk_tokens
        self.max_table_tokens = max_table_tokens  # Prevent oversized table chunks
        self.target_tokens = int(max_tokens * target_utilization)
        self.encoding_name = encoding
        
        try:
            self.tokenizer = tiktoken.get_encoding(self.encoding_name)
        except ValueError:
            print(f"Warning: Unknown encoding '{self.encoding_name}', falling back to cl100k_base")
            self.tokenizer = tiktoken.get_encoding('cl100k_base')
            self.encoding_name = 'cl100k_base'
        
        self.separators = [
            "\n\n\n",      # Multiple newlines (major sections)
            "\n\n",        # Paragraph breaks
            "\n",          # Line breaks
            ". ",          # Sentence endings
            "! ",          # Exclamations
            "? ",          # Questions
            "; ",          # Semicolons
            ", ",          # Commas
            " ",           # Spaces
            ""             # Character level (last resort)
        ]
        
        print(f"RecursiveChunker initialized for Weaviate hybrid search:")
        print(f"  Max tokens: {self.max_tokens}")
        print(f"  Target tokens: {self.target_tokens}")
        print(f"  Min tokens: {self.min_chunk_tokens}")
        print(f"  Max table tokens: {self.max_table_tokens}")
        print(f"  Overlap: {self.overlap_tokens}")
        print(f"  Optimized for text-embedding-3-small (1536 dimensions)")
    
    def count_tokens(self, text: str) -> int:
        return len(self.tokenizer.encode(text))
    
    def clean_html_artifacts(self, text: str) -> str:
        """Clean HTML artifacts and malformed markup"""
        # Remove <br> tags
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        
        # Remove other HTML tags but keep content
        text = re.sub(r'<[^>]+>', '', text)
        
        # Clean up malformed table separators
        text = re.sub(r'\|\s*-+\s*\|', '|---|', text)  # Normalize table separators
        text = re.sub(r'^-+\s*\|\s*-+.*$', '', text, flags=re.MULTILINE)  # Remove broken separators
        
        # Remove standalone separator lines that aren't part of tables
        text = re.sub(r'^[-|:\s]+$', '', text, flags=re.MULTILINE)
        
        return text
    
    def detect_real_table(self, text: str) -> Tuple[bool, int, int]:
        """
        Detect actual tables vs scattered pipe characters
        Returns: (is_table, start_line, end_line)
        """
        lines = text.split('\n')
        table_lines = []
        
        for i, line in enumerate(lines):
            # Count pipes, but exclude URLs and navigation
            if ('|' in line and 
                not re.search(r'https?://', line) and 
                not re.search(r'\[(URL|Home|Log\s+In)\]', line) and
                line.count('|') >= 2):  # At least 2 pipes for table row
                table_lines.append(i)
        
        if len(table_lines) < 2:  # Need at least 2 rows for a table
            return False, -1, -1
        
        # Check for consecutive or near-consecutive table lines
        consecutive_groups = []
        current_group = [table_lines[0]]
        
        for i in range(1, len(table_lines)):
            if table_lines[i] - table_lines[i-1] <= 2:  # Allow 1 empty line gap
                current_group.append(table_lines[i])
            else:
                if len(current_group) >= 2:
                    consecutive_groups.append(current_group)
                current_group = [table_lines[i]]
        
        if len(current_group) >= 2:
            consecutive_groups.append(current_group)
        
        # Find the largest table group
        if consecutive_groups:
            largest_group = max(consecutive_groups, key=len)
            return True, largest_group[0], largest_group[-1]
        
        return False, -1, -1
    
    def extract_table_section(self, text: str, start_line: int, end_line: int) -> str:
        """Extract table with some context"""
        lines = text.split('\n')
        
        # Include a bit of context before and after
        context_start = max(0, start_line - 1)
        context_end = min(len(lines), end_line + 2)
        
        return '\n'.join(lines[context_start:context_end])
    
    def detect_table(self, text: str) -> bool:
        """Improved table detection"""
        is_table, _, _ = self.detect_real_table(text)
        return is_table
    
    def detect_code_block(self, text: str) -> bool:
        """Detect code blocks"""
        # Markdown code fences
        fenced_code = bool(re.search(r'```.*?```', text, re.DOTALL))
        
        # Indented code blocks (4+ spaces or tab)
        indented_lines = len(re.findall(r'^(    |\t).+', text, re.MULTILINE))
        indented_code = indented_lines >= 3
        
        return fenced_code or indented_code
    
    def detect_list_structure(self, text: str) -> bool:
        """Detect list structures that should stay together"""
        # Numbered lists
        numbered = len(re.findall(r'^\d+\.\s', text, re.MULTILINE)) >= 3
        
        # Bullet lists
        bulleted = len(re.findall(r'^[-*+]\s', text, re.MULTILINE)) >= 3
        
        # Nested lists
        nested = bool(re.search(r'^(\s{2,}|\t)[-*+\d]', text, re.MULTILINE))
        
        return numbered or bulleted or nested
    
    def clean_text(self, text: str) -> str:
        """Clean extracted text while preserving structure"""
        # First clean HTML artifacts
        text = self.clean_html_artifacts(text)
        
        # Remove URLs but keep structure
        text = re.sub(r'https?://[^\s\]]+', '[URL]', text)
        
        # Remove navigation elements but preserve content structure
        nav_pattern = r'\[(Home|Ecosystem|Facilities|Management|Architecture|Marketplace|About us|Search\.\.\.|Log In|Ask AI)\]'
        text = re.sub(nav_pattern, '', text)
        
        # Fix spacing issues while preserving table structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            if '|' in line and line.count('|') >= 2:  # Preserve table formatting
                # Clean up table line but keep structure
                cleaned_line = re.sub(r'\s+', ' ', line.strip())
                cleaned_lines.append(cleaned_line)
            else:
                # Clean non-table lines
                line = re.sub(r'([a-z])([A-Z])', r'\1 \2', line)
                line = re.sub(r'[ \t]+', ' ', line)
                cleaned_lines.append(line.strip())
        
        text = '\n'.join(cleaned_lines)
        
        # Normalize whitespace but preserve structure
        text = re.sub(r'\n{4,}', '\n\n\n', text)  # Max 3 newlines
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove empty lines that aren't meaningful
        lines = text.split('\n')
        filtered_lines = []
        for i, line in enumerate(lines):
            if line.strip() or (i > 0 and i < len(lines) - 1):  # Keep structure
                filtered_lines.append(line)
        
        return '\n'.join(filtered_lines).strip()
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pymupdf4llm with better cleaning"""
        try:
            text = pymupdf4llm.to_markdown(str(file_path))
            print(f"  Extracted using pymupdf4llm (table-aware)")
            return text
        except Exception as e:
            print(f"  pymupdf4llm extraction failed: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX with structure preservation"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        if level.isdigit():
                            hashes = '#' * int(level)
                        else:
                            hashes = '##'
                        text_parts.append(f"\n\n{hashes} {paragraph.text.strip()}\n")
                    else:
                        text_parts.append(paragraph.text.strip())
            
            # Handle tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text.append(f"| {row_text} |")
                
                if table_text:
                    text_parts.append("\n\n" + "\n".join(table_text) + "\n")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        print(f"Warning: Could not decode {file_path} with common encodings")
        return ""
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract and clean text from file"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                raw_text = self.extract_text_from_pdf(file_path)
            elif extension == '.docx':
                raw_text = self.extract_text_from_docx(file_path)
            elif extension == '.txt':
                raw_text = self.extract_text_from_txt(file_path)
            else:
                print(f"Unsupported file type: {extension}")
                return ""
            
            return self.clean_text(raw_text)
        except Exception as e:
            print(f"Error extracting from {file_path.name}: {e}")
            return ""
    
    def split_large_table(self, text: str) -> List[str]:
        """Split large tables at natural boundaries"""
        lines = text.split('\n')
        sections = []
        current_section = []
        current_tokens = 0
        
        for line in lines:
            line_tokens = self.count_tokens(line)
            
            # If adding this line would exceed table limit, start new section
            if current_tokens + line_tokens > self.max_table_tokens and current_section:
                sections.append('\n'.join(current_section))
                current_section = [line]
                current_tokens = line_tokens
            else:
                current_section.append(line)
                current_tokens += line_tokens
        
        if current_section:
            sections.append('\n'.join(current_section))
        
        return sections if sections else [text]
    
    def structure_aware_pre_split(self, text: str) -> List[str]:
        """Split text while handling tables intelligently"""
        sections = []
        
        # First, identify all table regions
        is_table, start_line, end_line = self.detect_real_table(text)
        
        if is_table:
            lines = text.split('\n')
            
            # Split around the table
            before_table = '\n'.join(lines[:start_line]).strip()
            table_content = '\n'.join(lines[start_line:end_line + 1]).strip()
            after_table = '\n'.join(lines[end_line + 1:]).strip()
            
            # Add before-table content
            if before_table and self.count_tokens(before_table) >= self.min_chunk_tokens:
                sections.append(before_table)
            
            # Handle table content
            table_tokens = self.count_tokens(table_content)
            if table_tokens > self.max_table_tokens:
                # Split large table
                table_sections = self.split_large_table(table_content)
                sections.extend(table_sections)
            else:
                sections.append(table_content)
            
            # Add after-table content
            if after_table and self.count_tokens(after_table) >= self.min_chunk_tokens:
                sections.append(after_table)
            elif after_table and sections:
                # Merge small after-content with last section if possible
                last_section = sections[-1]
                combined = last_section + '\n\n' + after_table
                if self.count_tokens(combined) <= self.max_tokens:
                    sections[-1] = combined
        else:
            # No tables detected, use header-based splitting
            sections = self.header_based_split(text)
        
        return sections if sections else [text]
    
    def header_based_split(self, text: str) -> List[str]:
        """Split by headers while targeting optimal chunk sizes"""
        sections = []
        
        # Split by major headers (# and ##)
        if re.search(r'^#{1,3}\s+', text, re.MULTILINE):
            parts = re.split(r'^(#{1,3}\s+.+)$', text, flags=re.MULTILINE)
            current_section = ""
            
            for part in parts:
                if re.match(r'^#{1,3}\s+', part):
                    if current_section.strip():
                        token_count = self.count_tokens(current_section.strip())
                        if token_count >= self.min_chunk_tokens:
                            sections.append(current_section.strip())
                        elif sections:
                            # Try to merge with previous section
                            combined = sections[-1] + '\n\n' + current_section.strip()
                            if self.count_tokens(combined) <= self.max_tokens:
                                sections[-1] = combined
                            else:
                                sections.append(current_section.strip())
                    current_section = part
                else:
                    current_section += part
            
            if current_section.strip():
                sections.append(current_section.strip())
        
        # If no headers, build optimal-sized sections
        if not sections:
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            current_section = ""
            
            for para in paragraphs:
                test_section = current_section + '\n\n' + para if current_section else para
                test_tokens = self.count_tokens(test_section)
                
                if test_tokens <= self.target_tokens:
                    current_section = test_section
                else:
                    if current_section and self.count_tokens(current_section) >= self.min_chunk_tokens:
                        sections.append(current_section)
                    current_section = para
            
            if current_section and self.count_tokens(current_section) >= self.min_chunk_tokens:
                sections.append(current_section)
        
        return sections if sections else [text]
    
    def recursive_chunk(self, text: str, separators: List[str] = None, metadata: Dict = None) -> List[Dict]:
        if separators is None:
            separators = self.separators.copy()
        if metadata is None:
            metadata = {}
            
        text = text.strip()
        if not text:
            return []
        
        token_count = self.count_tokens(text)
        
        # Accept chunks in optimal range
        if self.min_chunk_tokens <= token_count <= self.max_tokens:
            quality_indicators = {
                'has_table': self.detect_table(text),
                'has_code': self.detect_code_block(text),
                'has_list': self.detect_list_structure(text),
                'good_length': token_count >= self.target_tokens * 0.7
            }
            
            return [{
                'text': text,
                'tokens': token_count,
                'metadata': {
                    **metadata, 
                    'chunk_method': 'semantic_optimal',
                    'quality_indicators': quality_indicators
                }
            }]
        
        # Special handling for structures, but with size limits
        if (self.detect_table(text) or self.detect_code_block(text) or 
            self.detect_list_structure(text)):
            
            if token_count <= self.max_table_tokens:
                # Preserve structure if reasonably sized
                structure_type = ('table' if self.detect_table(text) else 
                                'code' if self.detect_code_block(text) else 'list')
                return [{
                    'text': text,
                    'tokens': token_count,
                    'metadata': {
                        **metadata, 
                        'chunk_method': 'structure_preserved',
                        'structure_type': structure_type
                    }
                }]
            else:
                # Force split large structures
                if self.detect_table(text):
                    table_sections = self.split_large_table(text)
                    chunks = []
                    for section in table_sections:
                        chunks.extend(self.recursive_chunk(section, separators, metadata))
                    return chunks
        
        # Regular recursive splitting
        for sep_idx, separator in enumerate(separators):
            if separator in text or separator == "":
                if separator == "":
                    return self._character_level_split(text, metadata)
                
                parts = text.split(separator)
                chunks = []
                current_chunk = ""
                
                for i, part in enumerate(parts):
                    if i > 0:
                        part_with_sep = separator + part
                    else:
                        part_with_sep = part
                    
                    test_chunk = current_chunk + part_with_sep
                    test_tokens = self.count_tokens(test_chunk)
                    
                    if test_tokens <= self.target_tokens:
                        current_chunk = test_chunk
                    elif test_tokens <= self.max_tokens and self.count_tokens(current_chunk) < self.target_tokens * 0.5:
                        current_chunk = test_chunk
                    else:
                        if current_chunk.strip() and self.count_tokens(current_chunk.strip()) >= self.min_chunk_tokens:
                            chunks.append({
                                'text': current_chunk.strip(),
                                'tokens': self.count_tokens(current_chunk.strip()),
                                'metadata': {**metadata, 'chunk_method': f'split_by_{repr(separator)}'}
                            })
                        
                        if self.count_tokens(part_with_sep) > self.max_tokens:
                            remaining_separators = separators[sep_idx+1:]
                            sub_chunks = self.recursive_chunk(part_with_sep, remaining_separators, metadata)
                            chunks.extend(sub_chunks)
                            current_chunk = ""
                        else:
                            if chunks and self.overlap_tokens > 0:
                                overlap_text = self.get_semantic_overlap(chunks[-1]['text'])
                                current_chunk = overlap_text + part_with_sep
                            else:
                                current_chunk = part_with_sep
                
                if current_chunk.strip():
                    chunk_tokens = self.count_tokens(current_chunk.strip())
                    chunks.append({
                        'text': current_chunk.strip(),
                        'tokens': chunk_tokens,
                        'metadata': {**metadata, 'chunk_method': f'split_by_{repr(separator)}_final'}
                    })
                
                if chunks:
                    return chunks
        
        return [{
            'text': text,
            'tokens': token_count,
            'metadata': {**metadata, 'chunk_method': 'fallback'}
        }]
    
    def get_semantic_overlap(self, text: str) -> str:
        """Get semantically meaningful overlap"""
        sentences = re.split(r'[.!?]+\s+', text)
        
        if len(sentences) <= 1:
            tokens = self.tokenizer.encode(text)
            if len(tokens) <= self.overlap_tokens:
                return text + " "
            
            overlap_tokens = tokens[-self.overlap_tokens:]
            return self.tokenizer.decode(overlap_tokens) + " "
        
        overlap_text = sentences[-1]
        if len(sentences) > 1 and self.count_tokens(sentences[-2] + ". " + overlap_text) <= self.overlap_tokens:
            overlap_text = sentences[-2] + ". " + overlap_text
        
        return overlap_text + " "
    
    def _character_level_split(self, text: str, metadata: Dict) -> List[Dict]:
        """Last resort character-level splitting"""
        chunks = []
        start = 0
        
        while start < len(text):
            estimated_end = start + (self.target_tokens * 4)
            chunk = text[start:min(estimated_end, len(text))]
            
            while self.count_tokens(chunk) > self.max_tokens and len(chunk) > 1:
                chunk = chunk[:-1]
            
            if chunk:
                chunks.append({
                    'text': chunk,
                    'tokens': self.count_tokens(chunk),
                    'metadata': {**metadata, 'chunk_method': 'character_split'}
                })
            
            overlap_chars = int(len(chunk) * (self.overlap_tokens / self.max_tokens))
            start = start + len(chunk) - overlap_chars
            
            if start >= len(text):
                break
        
        return chunks
    
    def process_file(self, file_path: str) -> List[Dict]:
        """Process single file with improved table handling"""
        print(f"Processing: {Path(file_path).name}")
        
        raw_text = self.extract_text_from_file(file_path)
        if not raw_text:
            print(f"  No text extracted from {Path(file_path).name}")
            return []
        
        sections = self.structure_aware_pre_split(raw_text)
        print(f"  Split into {len(sections)} semantic/structural sections")
        
        all_chunks = []
        for section_idx, section in enumerate(sections):
            if not section.strip():
                continue
                
            metadata = {
                'source': Path(file_path).name,
                'file_path': str(file_path),
                'section_id': section_idx,
                'total_sections': len(sections),
                'embedding_model': 'text-embedding-3-small',
                'embedding_dimensions': 1536
            }
            
            section_chunks = self.recursive_chunk(section, metadata=metadata)
            all_chunks.extend(section_chunks)
        
        # Add final metadata
        for chunk_idx, chunk in enumerate(all_chunks):
            chunk['metadata']['chunk_id'] = chunk_idx
            chunk['metadata']['global_position'] = chunk_idx / len(all_chunks)
            
            if chunk_idx > 0:
                chunk['metadata']['prev_chunk_id'] = chunk_idx - 1
            if chunk_idx < len(all_chunks) - 1:
                chunk['metadata']['next_chunk_id'] = chunk_idx + 1
        
        print(f"  Created {len(all_chunks)} semantic chunks")
        return all_chunks
    
    def analyze_chunk_quality(self, chunks: List[Dict]) -> Dict:
        """Comprehensive quality analysis"""
        if not chunks:
            return {'error': 'No chunks to analyze'}
        
        token_counts = [chunk['tokens'] for chunk in chunks]
        methods = {}
        structure_preserved = 0
        oversized_chunks = 0
        
        for chunk in chunks:
            method = chunk['metadata'].get('chunk_method', 'unknown')
            methods[method] = methods.get(method, 0) + 1
            
            if 'structure_preserved' in method or 'semantic_optimal' in method:
                structure_preserved += 1
            
            if chunk['tokens'] > self.max_tokens:
                oversized_chunks += 1
        
        avg_tokens = sum(token_counts) / len(token_counts)
        utilization = (avg_tokens / self.max_tokens) * 100
        
        optimal_range = sum(1 for t in token_counts if self.target_tokens * 0.8 <= t <= self.max_tokens)
        small_chunks = sum(1 for t in token_counts if t < self.min_chunk_tokens)
        
        # Improved quality score
        semantic_score = (
            (optimal_range / len(chunks)) * 40 +
            (structure_preserved / len(chunks)) * 25 +
            (utilization / 100) * 20 +
            (1 - oversized_chunks / len(chunks)) * 15  # Penalty for oversized chunks
        ) * 100
        
        return {
            'total_chunks': len(chunks),
            'avg_tokens_per_chunk': round(avg_tokens, 1),
            'token_utilization': round(utilization, 1),
            'optimal_range_percent': round(optimal_range / len(chunks) * 100, 1),
            'structure_preserved_percent': round(structure_preserved / len(chunks) * 100, 1),
            'small_chunks_count': small_chunks,
            'oversized_chunks_count': oversized_chunks,
            'chunk_methods_used': methods,
            'files_processed': len(set(chunk['metadata']['source'] for chunk in chunks)),
            'semantic_quality_score': round(semantic_score, 1),
            'target_tokens': self.target_tokens,
            'min_tokens': self.min_chunk_tokens,
            'max_tokens': self.max_tokens,
            'max_table_tokens': self.max_table_tokens,
            'weaviate_ready': oversized_chunks == 0,
            'embedding_model': 'text-embedding-3-small'
        }
